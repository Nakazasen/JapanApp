import re
import os
import sys

# Add project root to path so we can import frontend
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from docx import Document
from sqlmodel import Session, select
from frontend.core.database import engine
from frontend.models.kanji import KanjiDeck, KanjiItem, KanjiVocab

# Regex patterns
# Group 1: Kanji, Group 2: Hán Việt, Group 3: Nghĩa (Optional)
KANJI_HEADER_PATTERN = r"^([一-龠])\s+([a-zA-ZĂăÂâĐđÊêÔôƠơƯưÀ-Ỹà-ỹ\s]+)\s*(?:\((.*)\))?"

# Group 1: Word, Group 2: Reading, Group 3: Hán Việt, Group 4: Nghĩa
VOCAB_PATTERN = r"^\d+\s+([^\s]+)\s+([ぁ-んァ-ヶ]+)\s+([A-ZĐÀ-Ỹa-zA-ZĂăÂâĐđÊêÔôƠơƯưÀ-Ỹà-ỹ\s]+)\s+(.+)"

# Group 1: Level (N1-N5)
LEVEL_HEADER_PATTERN = r"(?:JLPT\s+|Level\s+)?(N[1-5])\b"

def import_docx_to_db(docx_path: str, db_engine, default_deck_name="Imported Kanji", limit_paragraphs: int = None):
    # Ensure all tables exist
    from sqlmodel import SQLModel
    SQLModel.metadata.create_all(db_engine)

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"File not found: {docx_path}")

    doc = Document(docx_path)
    count_kanji = 0
    count_vocab = 0
    
    current_kanji_item = None
    current_level = None
    
    # Apply limit if provided
    paragraphs = doc.paragraphs[:limit_paragraphs] if limit_paragraphs else doc.paragraphs

    with Session(db_engine) as session:
        # Default deck
        current_deck = KanjiDeck(name=default_deck_name, description="Imported from Word", icon="📚")
        session.add(current_deck)
        session.commit()
        session.refresh(current_deck)
        
        # 2. Process paragraphs
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Check for Level Header
            level_match = re.search(LEVEL_HEADER_PATTERN, text)
            if level_match:
                # To distinguish section headers from TOC entries at start:
                # TOC entries often have page numbers or extra text.
                # Section headers are usually just the level or very short.
                is_header = len(text) < 30 or "合成漢字" in text
                
                if is_header:
                    new_level = level_match.group(1).upper()
                    if new_level != current_level:
                        current_level = new_level
                        deck_name = f"JLPT {current_level} 漢字"
                        
                        # Switch or create deck
                        statement = select(KanjiDeck).where(KanjiDeck.name == deck_name)
                        level_deck = session.exec(statement).first()
                        if not level_deck:
                            level_deck = KanjiDeck(
                                name=deck_name, 
                                description=f"Bộ thẻ {current_level} nhập từ Word", 
                                icon="🔖"
                            )
                            session.add(level_deck)
                            session.commit()
                            session.refresh(level_deck)
                        current_deck = level_deck
                        print(f"--- Switching to Deck: {deck_name} (at P{i}: {text}) ---")
                        continue

            # Check if it's a Kanji Header
            kanji_match = re.match(KANJI_HEADER_PATTERN, text)
            if kanji_match:
                char = kanji_match.group(1)
                hv = kanji_match.group(2).strip().upper()
                meaning = kanji_match.group(3).strip() if kanji_match.group(3) else ""
                
                # Check if kanji already exists
                existing = session.exec(select(KanjiItem).where(KanjiItem.kanji == char)).first()
                if existing:
                    existing.han_viet = hv
                    existing.meaning_vi = meaning
                    existing.jlpt_level = current_level
                    existing.deck_id = current_deck.id # Update deck if we found a level header
                    current_kanji_item = existing
                    session.add(existing)
                else:
                    current_kanji_item = KanjiItem(
                        kanji=char,
                        han_viet=hv,
                        meaning_vi=meaning,
                        jlpt_level=current_level,
                        deck_id=current_deck.id
                    )
                    session.add(current_kanji_item)
                    count_kanji += 1
                session.commit()
                continue

            # Check if it's a Vocab line
            vocab_match = re.match(VOCAB_PATTERN, text)
            if vocab_match and current_kanji_item:
                word = vocab_match.group(1).strip()
                # Avoid duplicate vocab for the same kanji
                existing_vocab = session.exec(
                    select(KanjiVocab).where(KanjiVocab.kanji_id == current_kanji_item.id, KanjiVocab.word == word)
                ).first()
                
                if not existing_vocab:
                    vocab_item = KanjiVocab(
                        word=word,
                        reading=vocab_match.group(2).strip(),
                        han_viet=vocab_match.group(3).strip(),
                        meaning_vi=vocab_match.group(4).strip(),
                        kanji_id=current_kanji_item.id
                    )
                    session.add(vocab_item)
                    count_vocab += 1
        
        session.commit()
        
    return count_kanji, count_vocab

if __name__ == "__main__":
    docx_file = r"D:\giao trinh day tieng nhat\videoHantu\dai tu dien kanji.docx"
    
    print(f"Starting FULL import from {docx_file}...")
    try:
        k, v = import_docx_to_db(docx_file, engine)
        print(f"Success! Added {k} Kanji and {v} vocab items.")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
